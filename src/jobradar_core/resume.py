from __future__ import annotations

import hashlib
import json
import re
import uuid
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from jobradar_core.database import LocalDatabase
from jobradar_core.models import (
    MemoryRecord,
    ResumeBlock,
    ResumeDiagnosis,
    ResumeDocument,
    ResumePatch,
    utc_now_iso,
)
from jobradar_core.workspace import Workspace, sha256_file

_KNOWN_KEYWORDS = (
    "Agent",
    "RAG",
    "LLM",
    "Python",
    "Java",
    "Go",
    "Golang",
    "C++",
    "SQL",
    "MySQL",
    "PostgreSQL",
    "Redis",
    "MongoDB",
    "Elasticsearch",
    "FastAPI",
    "Django",
    "Spring Boot",
    "Docker",
    "Kubernetes",
    "HTTP",
    "WebSocket",
    "RPC",
    "消息队列",
    "任务编排",
    "工具调用",
    "记忆管理",
    "上下文管理",
    "数据库",
    "分布式系统",
    "可观测性",
    "监控告警",
    "限流",
    "熔断",
    "重试",
    "降级",
    "CI/CD",
    "跨团队",
)
_NUMBER_RE = re.compile(r"(?<![A-Za-z])\d+(?:\.\d+)?%?")


class ResumeRepository:
    def __init__(self, database: LocalDatabase):
        self.database = database

    def save(self, document: ResumeDocument) -> None:
        with self.database.transaction() as connection:
            connection.execute(
                """INSERT OR IGNORE INTO resumes
                   (resume_id, original_name, source_path, source_hash, text,
                    blocks_json, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (
                    document.resume_id,
                    document.original_name,
                    document.source_path,
                    document.source_hash,
                    document.text,
                    json.dumps(
                        [block.model_dump() for block in document.blocks], ensure_ascii=False
                    ),
                    document.created_at,
                ),
            )

    def get(self, resume_id: str) -> ResumeDocument | None:
        with self.database.connect() as connection:
            row = connection.execute(
                "SELECT * FROM resumes WHERE resume_id=?", (resume_id,)
            ).fetchone()
        if not row:
            return None
        return ResumeDocument(
            resume_id=row["resume_id"],
            original_name=row["original_name"],
            source_path=row["source_path"],
            source_hash=row["source_hash"],
            text=row["text"],
            blocks=[ResumeBlock.model_validate(item) for item in json.loads(row["blocks_json"])],
            created_at=row["created_at"],
        )

    def find_by_hash(self, source_hash: str) -> ResumeDocument | None:
        with self.database.connect() as connection:
            row = connection.execute(
                "SELECT resume_id FROM resumes WHERE source_hash=?", (source_hash,)
            ).fetchone()
        return self.get(row["resume_id"]) if row else None

    def save_version(
        self,
        *,
        resume_id: str,
        run_id: str,
        path: Path,
        content_hash: str,
        patches: list[ResumePatch],
    ) -> str:
        version_id = uuid.uuid4().hex
        with self.database.transaction() as connection:
            connection.execute(
                """INSERT INTO resume_versions
                   (version_id, resume_id, run_id, path, content_hash, patches_json, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (
                    version_id,
                    resume_id,
                    run_id,
                    str(path),
                    content_hash,
                    json.dumps([patch.model_dump() for patch in patches], ensure_ascii=False),
                    utc_now_iso(),
                ),
            )
        return version_id


class ResumeParser:
    supported_suffixes = {".pdf", ".docx", ".md", ".txt"}

    def __init__(
        self,
        workspace: Workspace,
        repository: ResumeRepository,
        database: LocalDatabase,
    ):
        self.workspace = workspace
        self.repository = repository
        self.database = database

    def parse(self, path: Path) -> ResumeDocument:
        source = path.expanduser().resolve(strict=True)
        if source.suffix.lower() not in self.supported_suffixes:
            raise ValueError("简历格式仅支持 PDF、DOCX、Markdown 和 TXT")
        source_hash = sha256_file(source)
        existing = self.repository.find_by_hash(source_hash)
        if existing:
            return existing
        preserved = self.workspace.import_user_file(source, self.workspace.original_resumes)
        extracted = self._extract(preserved)
        if not extracted.strip():
            raise ValueError("没有从简历中提取到文字；扫描版 PDF 请先进行 OCR")
        blocks = _make_blocks(extracted, source_hash)
        document = ResumeDocument(
            resume_id=uuid.uuid4().hex,
            original_name=source.name,
            source_path=str(preserved),
            source_hash=source_hash,
            text=extracted,
            blocks=blocks,
        )
        self.repository.save(document)
        self._index_memory(document)
        return document

    def _extract(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".md", ".txt"}:
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".docx":
            doc = Document(path)
            return "\n".join(
                paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()
            )
        if suffix == ".pdf":
            pages = []
            for page in PdfReader(str(path)).pages:
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(text.strip())
            return "\n\n".join(pages)
        raise ValueError(f"unsupported resume format: {suffix}")

    def _index_memory(self, document: ResumeDocument) -> None:
        self.database.add_memory(
            MemoryRecord(
                memory_id=f"resume-index-{document.resume_id}",
                scope="resume",
                level="L1",
                category="resume_index",
                summary=f"简历 {document.original_name}，共 {len(document.blocks)} 个证据块",
                payload={"resume_id": document.resume_id, "source_hash": document.source_hash},
                source_ref=f"resume:{document.resume_id}",
                user_confirmed=True,
                status="active",
            )
        )
        for block in document.blocks[:240]:
            self.database.add_memory(
                MemoryRecord(
                    memory_id=f"resume-evidence-{document.resume_id}-{block.block_id}",
                    scope="resume",
                    level="L3",
                    category="resume_evidence",
                    summary=block.text[:180],
                    payload={"resume_id": document.resume_id, "block_id": block.block_id},
                    source_ref=block.source_ref,
                    user_confirmed=True,
                    status="active",
                )
            )


def local_diagnosis(resume_text: str, jd_text: str) -> ResumeDiagnosis:
    resume_lower = resume_text.lower()
    jd_lower = jd_text.lower()
    jd_keywords = [keyword for keyword in _KNOWN_KEYWORDS if keyword.lower() in jd_lower]
    matched = [keyword for keyword in jd_keywords if keyword.lower() in resume_lower]
    missing = [keyword for keyword in jd_keywords if keyword.lower() not in resume_lower]
    strengths = [f"简历中已有「{keyword}」相关证据" for keyword in matched[:8]]
    gaps = [f"JD 提到「{keyword}」，简历中未找到直接证据" for keyword in missing[:8]]
    if not jd_text.strip():
        summary = "尚未提供目标 JD，只完成了简历结构化解析。"
    elif matched:
        summary = f"找到 {len(matched)} 个直接匹配关键词，另有 {len(missing)} 项需要核实。"
    else:
        summary = "当前简历与 JD 的直接关键词重合较少，应先核实可迁移经历再改写。"
    return ResumeDiagnosis(
        summary=summary,
        strengths=strengths,
        gaps=gaps,
        matched_keywords=matched,
        missing_keywords=missing,
    )


def verify_patches(
    patches: list[ResumePatch],
    document: ResumeDocument,
    diagnosis: ResumeDiagnosis,
) -> list[ResumePatch]:
    block_map = {block.block_id: block for block in document.blocks}
    source_refs = {block.source_ref for block in document.blocks}
    original_numbers = set(_NUMBER_RE.findall(document.text))
    missing_lower = {keyword.lower() for keyword in diagnosis.missing_keywords}
    verified: list[ResumePatch] = []
    for patch in patches[:12]:
        risks = list(patch.risk_flags)
        block = block_map.get(patch.target_block_id)
        if block is None:
            risks.append("目标段落不存在")
        elif patch.before != block.text:
            risks.append("before 与原简历不一致")
        if not patch.evidence_refs or any(ref not in source_refs for ref in patch.evidence_refs):
            risks.append("缺少有效的简历证据引用")
        elif block is not None and block.source_ref not in patch.evidence_refs:
            risks.append("证据引用未包含目标段落")
        if not patch.after.strip():
            risks.append("修改结果为空")
        elif patch.after == patch.before:
            risks.append("修改前后没有变化")
        introduced_numbers = set(_NUMBER_RE.findall(patch.after)) - original_numbers
        if introduced_numbers:
            risks.append(f"新增了无原文依据的数字：{'、'.join(sorted(introduced_numbers))}")
        introduced_missing = [
            keyword
            for keyword in missing_lower
            if keyword in patch.after.lower() and keyword not in patch.before.lower()
        ]
        if introduced_missing:
            risks.append("把 JD 缺口写成了候选人已具备事实")
        if patch.intent == "fact_needed":
            risks.append("需要用户补充或确认事实，不能直接应用")
        if risks:
            patch = patch.model_copy(update={"risk_flags": risks, "status": "blocked"})
        verified.append(patch)
    return verified


def export_resume_version(
    *,
    document: ResumeDocument,
    patches: list[ResumePatch],
    run_id: str,
    workspace: Workspace,
    repository: ResumeRepository,
) -> Path:
    accepted = [patch for patch in patches if patch.status in {"accepted", "edited"}]
    if not accepted:
        raise ValueError("没有已接受的修改可导出")
    output = document.text
    for patch in accepted:
        if patch.before not in output:
            raise ValueError(f"无法定位待替换原文：{patch.target_block_id}")
        output = output.replace(patch.before, patch.after, 1)
    timestamp = utc_now_iso().replace(":", "-")[:19]
    destination = workspace.resume_versions / f"{Path(document.original_name).stem}-{timestamp}.md"
    destination.write_text(output, encoding="utf-8")
    content_hash = hashlib.sha256(output.encode("utf-8")).hexdigest()
    repository.save_version(
        resume_id=document.resume_id,
        run_id=run_id,
        path=destination,
        content_hash=content_hash,
        patches=accepted,
    )
    return destination


def _make_blocks(text: str, source_hash: str) -> list[ResumeBlock]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return [
        ResumeBlock(
            block_id=f"b-{index:04d}",
            text=line,
            source_ref=f"resume:{source_hash[:12]}:line:{index}",
            order=index,
        )
        for index, line in enumerate(lines, start=1)
    ]
