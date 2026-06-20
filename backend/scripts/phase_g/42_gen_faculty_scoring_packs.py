"""生成"真人锚"打分批次 — 给 SAIF 老师/学生对系统推荐的岗位打分。

目的:把对口准确率从判官自评升级为真人锚。每个 persona 一张表:系统当前判为该赛道的
top-8 岗位(去重公司、GT 优先、近期),老师/学生逐条打两个分:
  对口?(0-3: 0完全不对口 / 1沾边 / 2基本对口 / 3非常对口)
  值得投?(0-3: 0不值得 / 1一般 / 2不错 / 3优质必投)
回收后这批分就是 nDCG 的金标准锚点(再用判官放大到大批量)。

跑:cd backend && PYTHONPATH=. .venv/bin/python scripts/phase_g/42_gen_faculty_scoring_packs.py
输出:data/_phase_g/scoring_packs/<persona>.md(再 cp 到 sync / 飞书给老师)
"""
from __future__ import annotations
import json, os, sqlite3

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _set_cjk(run, name="Microsoft YaHei", size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_fill})
    tcPr.append(shd)


def build_docx(pid, desc, track, picked, out_path):
    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    h = doc.add_paragraph()
    _set_cjk(h.add_run(f"推荐对口打分表 · {pid}"), size=16, bold=True)

    for label, val in (("学生画像", desc), ("系统判定赛道", track)):
        p = doc.add_paragraph()
        _set_cjk(p.add_run(f"{label}："), size=11, bold=True)
        _set_cjk(p.add_run(val), size=11)

    intro = (
        "下面 8 个岗位是系统当前判定为「该赛道」的代表岗位（每家公司取一条）。"
        "请您站在这位学生的角度，逐条打两个分（直接在右侧两列填数字）：\n"
        "· 对口?  0=完全不对口 / 1=沾边 / 2=基本对口 / 3=非常对口\n"
        "· 值得投?  0=不值得 / 1=一般 / 2=不错 / 3=优质必投\n"
        "末列「备注」可写一句为什么（尤其打 0/1 时，帮我们定位系统哪里判错）。"
    )
    pi = doc.add_paragraph()
    _set_cjk(pi.add_run(intro), size=10, color=(0x55, 0x55, 0x55))

    headers = ["#", "公司", "岗位", "系统判的细分方向", "对口?(0-3)", "值得投?(0-3)", "备注"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for j, htext in enumerate(headers):
        c = table.rows[0].cells[j]
        _shade(c, "C96442")
        run = c.paragraphs[0].add_run(htext)
        _set_cjk(run, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (co, title, sc, tier, isgt) in enumerate(picked, 1):
        cells = table.add_row().cells
        vals = [str(i), co + (" ⭐" if isgt else ""), title, sc, "", "", ""]
        for j, v in enumerate(vals):
            run = cells[j].paragraphs[0].add_run(v)
            _set_cjk(run, size=9)
            if j in (0, 4, 5):
                cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    s = doc.add_paragraph()
    _set_cjk(s.add_run("汇总（老师填）：这 8 条里，您认为真正对口（打 2-3 分）的有几条？ ____ / 8"),
             size=11, bold=True)
    f = doc.add_paragraph()
    _set_cjk(f.add_run("⭐ = 该公司在我们的金融重点公司名单内。"), size=9, color=(0x88, 0x88, 0x88))
    doc.save(out_path)

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DB = os.path.join(ROOT, "data", "jobradar.db")
REG = os.path.join(ROOT, "data", "_phase_g", "finance_taxonomy_v1.json")
GT = os.path.join(ROOT, "data", "ground_truth_companies_v1.json")
OUT = os.path.join(ROOT, "data", "_phase_g", "scoring_packs")

# persona → (展示名, 背景一句话, v1 track key)。只挑现池够 8 家公司的干净赛道。
PERSONAS = [
    ("P1-投研", "林思远 · 目标公募/资管二级市场行业研究员(应届)", "公募/资管·投研"),
    ("P2-卖方", "顾天瑜 · 目标券商研究所卖方研究员", "卖方研究"),
    ("P5-投行", "魏文骏 · 目标投行 IBD / 并购 / 资本市场", "投行·并购·资本市场"),
    ("P6-量化", "韩怀宇 · 目标量化研究 / 量化开发", "量化"),
]
PER = 8


def main():
    os.makedirs(OUT, exist_ok=True)
    reg = json.load(open(REG))
    track2subcats = {t["key"]: t["sub_cats"] for t in reg["tracks"]}
    # v1 名 → 库内实际字符串(DB 未归一)
    rev = {sc: [sc] for t in reg["tracks"] for sc in t["sub_cats"]}
    for old, new in reg["canonical_map_old_to_new"].items():
        if new in rev and old not in rev[new]:
            rev[new].append(old)
    gt = json.load(open(GT))
    gt_names = set()
    for lst in gt.get("ground_truth", {}).values():
        for e in lst:
            if e.get("name"):
                gt_names.add(e["name"])

    c = sqlite3.connect(DB).cursor()
    for pid, desc, track in PERSONAS:
        v1_scs = track2subcats[track]
        names = []
        for sc in v1_scs:
            names += rev.get(sc, [sc])
        names = list(set(names))
        ph = ",".join("?" * len(names))
        rows = c.execute(
            f"""SELECT company, job_title, sub_category, COALESCE(institution_tier,''), id
                FROM jobs WHERE sub_category IN ({ph})
                  AND quality_label IN('good','internship_only')
                  AND length(COALESCE(job_duty,'')||COALESCE(job_req,''))>=30
                ORDER BY id DESC""",
            names,
        ).fetchall()
        # 去重公司(每家取最新一条),GT 优先
        seen, picked = set(), []
        for co, title, sc, tier, jid in rows:
            if co in seen:
                continue
            seen.add(co)
            picked.append((co, title, sc, tier, co in gt_names))
        picked.sort(key=lambda x: (not x[4],))  # GT 公司排前
        picked = picked[:PER]

        lines = [
            f"# 推荐对口打分表 · {pid}",
            "",
            f"**学生画像**:{desc}",
            f"**系统判定赛道**:{track}",
            "",
            "> 下面 8 个岗位是系统当前判定为「该赛道」的代表岗位(每家公司取一条)。",
            "> 请您站在这位学生的角度,逐条打两个分(直接在表里填数字):",
            "> - **对口?** 0=完全不对口 / 1=沾边 / 2=基本对口 / 3=非常对口",
            "> - **值得投?** 0=不值得 / 1=一般 / 2=不错 / 3=优质必投",
            "> 末列「备注」可写一句为什么(尤其打 0/1 时,帮我们定位系统哪里判错)。",
            "",
            "| # | 公司 | 岗位 | 系统判的细分方向 | 对口?(0-3) | 值得投?(0-3) | 备注 |",
            "|---|---|---|---|:---:|:---:|---|",
        ]
        for i, (co, title, sc, tier, isgt) in enumerate(picked, 1):
            mark = " ⭐GT" if isgt else ""
            lines.append(f"| {i} | {co}{mark} | {title[:34]} | {sc} |  |  |  |")
        lines += [
            "",
            "---",
            "**汇总(老师填)**:这 8 条里,您认为**真正对口(打 2-3 分)的有几条?** ____ / 8",
            "",
            "⭐GT = 该公司在我们的金融重点公司名单内。",
        ]
        path = os.path.join(OUT, f"{pid}.md")
        open(path, "w").write("\n".join(lines))
        docx_path = os.path.join(OUT, f"{pid}.docx")
        build_docx(pid, desc, track, picked, docx_path)
        print(f"  {pid}: {len(picked)} 岗 → {path} + .docx")
    print(f"\n打分批次已生成(md + docx) → {OUT}")


if __name__ == "__main__":
    main()
